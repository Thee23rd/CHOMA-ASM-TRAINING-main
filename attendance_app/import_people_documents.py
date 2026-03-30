import math
import os
import re
from typing import Any, Dict, List, Optional

import pandas as pd

from db import DB_PATH, init_db, upsert_person, add_cooperative_placeholder


def _clean_str(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, float):
        if math.isnan(v):
            return ""
    s = str(v).strip()
    if s.lower() in {"nan", "none", ""}:
        return ""
    # Remove accidental double spaces.
    return " ".join(s.split())


def _normalize_contacts(v: Any) -> str:
    # Pandas often reads phone numbers as floats. Convert 123.0 -> "123".
    if v is None:
        return ""
    if isinstance(v, float):
        if math.isnan(v):
            return ""
        if v.is_integer():
            return str(int(v))
        return str(v).strip()
    return _clean_str(v)


def _map_gender(gender: Any) -> str:
    g = _clean_str(gender).upper()
    if g in {"F", "FEMALE"}:
        return "Female"
    if g in {"M", "MALE"}:
        return "Male"
    return "Other"


def import_from_xlsx(
    xlsx_path: str,
    sheet_prefer: Optional[str] = "COOPERATIVES TEMPLATE",
) -> Dict[str, int]:
    xlsx_path = os.path.abspath(xlsx_path)
    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(xlsx_path)

    xl = pd.ExcelFile(xlsx_path)
    sheet_name = sheet_prefer if sheet_prefer in xl.sheet_names else xl.sheet_names[0]

    df = pd.read_excel(xlsx_path, sheet_name=sheet_name, header=2)

    # Forward-fill cooperative/district because the template uses merged cells.
    df["NAME OF COOPERATIVE"] = df["NAME OF COOPERATIVE"].ffill()
    df["District and Village"] = df["District and Village"].ffill()

    inserted_or_updated = 0
    skipped = 0

    for _, r in df.iterrows():
        coop = _clean_str(r.get("NAME OF COOPERATIVE"))
        district = _clean_str(r.get("District and Village"))
        contact_name = _clean_str(r.get("Participants Names"))
        sex = _map_gender(r.get("Gender"))
        nrc = _clean_str(r.get("NRC"))
        contacts = _normalize_contacts(r.get("Contacts"))

        if not coop or not district:
            skipped += 1
            continue
        if contact_name and nrc:
            upsert_person(
                cooperative_name=coop,
                group_venue=None,
                district_location=district,
                contact_person_details=contact_name,
                sex=sex,
                nrc_details=nrc,
                contact_number=contacts or None,
            )
            inserted_or_updated += 1
        else:
            # Cooperative with no participants - add placeholder
            if add_cooperative_placeholder(coop, district):
                inserted_or_updated += 1
            else:
                skipped += 1

    return {"imported_or_updated": inserted_or_updated, "skipped": skipped, "placeholders": "included"}


def import_from_docx(docx_path: str) -> Dict[str, int]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX import. Run: pip install python-docx")

    docx_path = os.path.abspath(docx_path)
    if not os.path.exists(docx_path):
        raise FileNotFoundError(docx_path)

    doc = Document(docx_path)
    if not doc.tables:
        return {"imported_or_updated": 0, "skipped": 0}

    # The file seems to contain one relevant table with headers in the first row.
    table = doc.tables[0]
    header_row_index = 0
    num_cols = len(table.columns)
    headers = [table.cell(header_row_index, c).text.strip() for c in range(num_cols)]

    def _idx_contains(substr: str) -> Optional[int]:
        substr_low = substr.lower()
        for i, h in enumerate(headers):
            if substr_low in h.lower():
                return i
        return None

    coop_i = _idx_contains("NAME OF COOPERATIVE")
    district_i = _idx_contains("DISTRICT / LOCATION")
    contact_i = _idx_contains("CONTACT PERSON / DETAILS")
    nrc_i = _idx_contains("NRC")

    if coop_i is None or district_i is None or contact_i is None or nrc_i is None:
        raise RuntimeError(
            f"Could not map columns in docx. Headers found: {headers}"
        )

    imported = 0
    skipped = 0

    PHONE_RE = re.compile(r"\b\d{9,}\b")
    NRC_PAIR_RE = re.compile(r"([A-Za-z][A-Za-z .()]*?)\s*-\s*(\d+/\d+/\d+)")
    NRC_RE = re.compile(r"\d+/\d+/\d+")

    def extract_nrc_pairs(nrc_text: str) -> List[Dict[str, str]]:
        text = nrc_text or ""
        if not text.strip():
            return []

        pairs: List[Dict[str, str]] = []
        matches = list(NRC_PAIR_RE.finditer(text))
        for m in matches:
            name = _clean_str(m.group(1))
            nrc = _clean_str(m.group(2))
            if name and nrc:
                pairs.append({"name": name, "nrc": nrc})

        if pairs:
            return pairs

        # Fallback: extract NRCs and infer name from text preceding each NRC.
        nrc_matches = list(NRC_RE.finditer(text))
        last = 0
        for m in nrc_matches:
            before = text[last:m.start()].strip()
            before = before.rstrip("- ").strip()
            # If multiple participants exist, names are usually separated by 2 spaces.
            if "  " in before:
                before = before.split("  ")[-1]
            name = _clean_str(before)
            nrc = _clean_str(m.group(0))
            if name and nrc:
                pairs.append({"name": name, "nrc": nrc})
            last = m.end()
        return pairs

    def extract_phones(contact_text: str) -> List[Dict[str, Any]]:
        phone_matches = []
        for m in PHONE_RE.finditer(contact_text or ""):
            phone_matches.append({"phone": m.group(0), "pos": m.start()})
        return phone_matches

    def assign_phones_to_participants(contact_text: str, participants: List[Dict[str, str]]):
        phone_matches = extract_phones(contact_text)
        if not participants:
            return []
        if not phone_matches:
            return [None for _ in participants]

        assigned: List[Optional[str]] = [None for _ in participants]
        assigned_pos: List[Optional[int]] = [None for _ in participants]

        # Try name-based assignment first.
        lower_text = (contact_text or "").lower()
        for i, person in enumerate(participants):
            name = (person.get("name") or "").lower()
            if not name:
                continue
            idx = lower_text.find(name)
            if idx == -1:
                continue
            # Pick first phone that occurs after the name.
            for pm in phone_matches:
                if pm["pos"] > idx:
                    assigned[i] = pm["phone"]
                    assigned_pos[i] = pm["pos"]
                    break

        # Fill remaining using phone order (by position), skipping positions already used.
        used_positions = {p for p in assigned_pos if p is not None}
        phone_idx = 0
        for i in range(len(assigned)):
            if assigned[i] is not None:
                continue
            while phone_idx < len(phone_matches) and phone_matches[phone_idx]["pos"] in used_positions:
                phone_idx += 1
            if phone_idx >= len(phone_matches):
                break
            assigned[i] = phone_matches[phone_idx]["phone"]
            used_positions.add(phone_matches[phone_idx]["pos"])
            phone_idx += 1
        return assigned

    # Data starts after header row.
    for ri in range(header_row_index + 1, len(table.rows)):
        row = table.rows[ri]
        coop = _clean_str(row.cells[coop_i].text)
        district = _clean_str(row.cells[district_i].text)
        contact_text = _clean_str(row.cells[contact_i].text)
        nrc_text = _clean_str(row.cells[nrc_i].text)

        if not coop or not district or not nrc_text:
            skipped += 1
            continue

        participants = extract_nrc_pairs(nrc_text)
        if not participants:
            skipped += 1
            continue

        phones_assigned = assign_phones_to_participants(contact_text, participants)

        for idx, person in enumerate(participants):
            name = _clean_str(person.get("name"))
            nrc = _clean_str(person.get("nrc"))
            if not name or not nrc:
                skipped += 1
                continue

            upsert_person(
                cooperative_name=coop,
                group_venue=None,
                district_location=district,
                contact_person_details=name,
                sex="Other",
                nrc_details=nrc,
                contact_number=phones_assigned[idx],
            )
            imported += 1

    return {"imported_or_updated": imported, "skipped": skipped}


def import_from_final_register(xlsx_path: str) -> Dict[str, Any]:
    """
    Import from FINAL REGISTER 2.xlsx - both COOPERATIVES TEMPLATE and UPDATED REGISTER sheets.
    Uses same logic as import_from_xlsx (people + placeholders for cooperatives with no members).
    """
    xlsx_path = os.path.abspath(xlsx_path)
    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(xlsx_path)
    total = {"imported_or_updated": 0, "skipped": 0}
    xl = pd.ExcelFile(xlsx_path)
    sheets = [s for s in ["COOPERATIVES TEMPLATE", "UPDATED REGISTER"] if s in xl.sheet_names]
    if not sheets:
        sheets = xl.sheet_names[:2]
    for sheet_name in sheets:
        res = import_from_xlsx(xlsx_path, sheet_prefer=sheet_name)
        total["imported_or_updated"] += res["imported_or_updated"]
        total["skipped"] += res["skipped"]
    return total


def main() -> None:
    init_db(DB_PATH)

    workspace_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

    # These are the filenames you shared in your `CHOMA TRAINING` folder.
    xlsx = os.path.join(workspace_root, "UPND SECRETARIAT HQS FINAL LIST FOR  COORPERATIVES.xlsx")
    docx = os.path.join(workspace_root, "CHOMA ASM TRAINING LIST.docx")
    final_xlsx = os.path.join(workspace_root, "FINAL REGISTER 2.xlsx")

    print(f"DB_PATH: {DB_PATH}")
    print("Importing from:", xlsx)
    xlsx_res = import_from_xlsx(xlsx)
    print("XLSX result:", xlsx_res)

    print("Importing from:", docx)
    docx_res = import_from_docx(docx)
    print("DOCX result:", docx_res)

    if os.path.exists(final_xlsx):
        print("Importing from:", final_xlsx)
        final_res = import_from_final_register(final_xlsx)
        print("FINAL REGISTER result:", final_res)
    else:
        print("FINAL REGISTER 2.xlsx not found, skipping")

    print("Done.")


if __name__ == "__main__":
    main()

